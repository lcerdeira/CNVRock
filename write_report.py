"""Generate the CNVRock diagnostics report as a DOCX file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = "assets/screenshots"
OUT = "assets/CNVRock_diagnostics_report.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Styles helpers ────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

def caption(text):
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run()
    run.italic = True
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    return p

def figure(path, width_inches=5.5, cap=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        body(f"[Figure not found: {path}]")
    if cap:
        caption(cap)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p

def table_row(tbl, cells):
    row = tbl.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row

# ─────────────────────────────────────────────────────────────────────────────
# Title
# ─────────────────────────────────────────────────────────────────────────────
doc.add_heading("CNVRock — AMR Gene Copy-Number Detection in K. pneumoniae", 0)
p = doc.add_paragraph("Diagnostics report · Experiment 29 · May 2026")
p.runs[0].italic = True
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 1. Background
# ─────────────────────────────────────────────────────────────────────────────
heading("1. Background and motivation")

body(
    "Antimicrobial resistance (AMR) in Klebsiella pneumoniae Species Complex (KpSC) is "
    "driven not only by gene presence but also by copy number. A bacterium carrying multiple "
    "plasmid copies of blaKPC-2, for instance, expresses higher carbapenemase levels than one "
    "with a single copy, with direct implications for treatment failure. Similarly, chromosomal "
    "amplification of blaSHV or loss of porin genes (ompK35, ompK36) shifts resistance phenotype "
    "in ways that gene presence/absence calling alone cannot capture."
)
body(
    "CNVRock detects AMR gene copy-number variation (CNV) from whole-genome sequencing (WGS) "
    "read-depth data. It uses a convolutional variational autoencoder (VAE) to learn a compact "
    "latent representation of genome-wide depth, a Gaussian hidden Markov model (HMM) to segment "
    "latent trajectories into discrete copy-number states, and a two-track caller that handles "
    "both chromosomal and plasmid-borne genes differently."
)
body(
    "The 545-sample cohort is drawn from AllTheBacteria (KpSC assemblies, quality-filtered). "
    "Ground truth is derived from AMRFinder+ (plasmid genes) and Kleborate v3 (porin/efflux "
    "status). Experiment 29 is the current best-performing configuration."
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. Pipeline overview
# ─────────────────────────────────────────────────────────────────────────────
heading("2. Pipeline overview")

body(
    "The pipeline processes per-sample read-count vectors extracted from BAMs aligned to the "
    "HS11286 reference genome. Two parallel data tracks are maintained:"
)
bullet("Chromosomal track — 1 kb bins across the core genome (~5 334 bins per sample). "
       "Passed through the VAE + HMM to detect blaSHV amplification and ramR deletion.")
bullet("Plasmid track — unmapped reads remapped individually to one plasmid reference per gene "
       "(blaKPC-2, blaCTX-M-15, blaNDM-1, qnrB1, blaOXA-48, aac6-Ib-cr, blaCTX-M-14, blaTEM-1). "
       "Plasmid copy number (PCN) is computed directly as mean gene depth / chromosomal median.")

body("The key processing steps are:")
bullet("Step 1 — VAE training: a 10-dimensional latent space is learnt from normalised "
       "read-depth profiles. β-VAE with KL warmup forces the latent to encode only "
       "biologically meaningful depth variation.")
bullet("Step 2 — HMM segmentation: a 6-state Gaussian HMM segments each sample's latent "
       "trajectory along the chromosome into copy-number segments (CN 0–5).")
bullet("Step 3 — Chromosomal gene calling: copy ratios (gene depth / flanking depth) gate "
       "the HMM calls into per-gene CN integers.")
bullet("Step 4 — Plasmid gene calling: PCN thresholds (per-gene, from plasmid_gene_coords.tsv) "
       "classify each gene as absent (CN 0), present (CN 1), or amplified (CN 2).")
bullet("Step 5 — Evaluation: MCC, FNR, and PPV computed against AMRFinder / Kleborate ground truth.")

# ─────────────────────────────────────────────────────────────────────────────
# 3. Experiment 29 results summary
# ─────────────────────────────────────────────────────────────────────────────
heading("3. Experiment 29 — performance summary")

body(
    "The table below summarises classification performance across all evaluated genes. "
    "FNR (false-negative rate) is the primary optimisation target: a missed carbapenemase "
    "carrier could receive inappropriate treatment. PPV (positive predictive value) is "
    "secondary — apparent false positives may reflect AMRFinder's incomplete variant coverage."
)

tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Gene", "Type", "MCC", "FNR", "PPV", "Notes"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

rows = [
    ("blaKPC-2",    "plasmid",       "1.00", "0.00", "1.00", "Perfect calling"),
    ("blaCTX-M-15", "plasmid",       "0.82", "0.20", "1.00", "53 irreducible FNs; blaSHV cross-map"),
    ("blaNDM-1",    "plasmid",       "0.99", "0.00", "0.99", ""),
    ("qnrB1",       "plasmid",       "0.98", "0.03", "1.00", "GT expanded to all qnrB variants"),
    ("blaOXA-48",   "plasmid",       "0.98", "0.01", "0.99", "GT expanded to OXA-48-like family"),
    ("aac6-Ib-cr",  "plasmid",       "0.86", "0.13", "0.93", "Low PCN; threshold lowered to 0.10"),
    ("blaSHV",      "chromosomal",   "—",    "—",    "—",    "Calling in progress"),
    ("ramR",        "chromosomal",   "—",    "—",    "—",    "Calling in progress"),
]
for r in rows:
    table_row(tbl, r)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 4. Figures
# ─────────────────────────────────────────────────────────────────────────────
heading("4. Diagnostic figures")

# 4.1 PCA
heading("4.1  Latent space PCA", level=2)

body(
    "The VAE compresses each sample's genome-wide depth profile into a 10-dimensional latent "
    "vector. Principal component analysis (PCA) of those latent vectors is shown below, with "
    "PC1 explaining 86.2% and PC2 explaining 9.7% of total variance."
)
figure(
    f"{SCREENSHOTS}/exp29_pca_latents.png",
    width_inches=4.5,
    cap="Figure 1. PCA of VAE latent vectors for 545 KpSC samples (experiment 29). "
        "Each dot represents one sample. The highlighted sample (ERR10461880, red) is "
        "blaKPC-2 positive. The two clusters along PC1 broadly separate samples with "
        "elevated overall depth (right, positive PC1) from those with more typical profiles "
        "(left, negative PC1). The lone outlier at PC2 ≈ 19 represents an unusual depth profile "
        "warranting manual inspection."
)
body(
    "Why this matters: the PCA provides a global quality-control view. Samples that "
    "cluster far from the main population may indicate coverage anomalies, contamination, "
    "or genuinely unusual genomic structure. Colour-coding by gene call status (available "
    "interactively in the app) reveals whether the latent organisation correlates with "
    "resistance gene content."
)

# 4.2 Latent bar
heading("4.2  Per-sample latent vector", level=2)

body(
    "Each of the 10 latent dimensions captures a distinct axis of variation in the "
    "read-depth signal. The bar chart below shows the latent vector for sample ERR10461880."
)
figure(
    f"{SCREENSHOTS}/exp29_latent_bar_ERR10461880.png",
    width_inches=5.5,
    cap="Figure 2. Latent vector for sample ERR10461880 (blaKPC-2 positive). "
        "Bar height represents the value along each latent dimension (z1–z10); "
        "colour intensity reflects absolute magnitude. z9 is the most strongly "
        "negative dimension for this sample; z1 and z3 are the most positive. "
        "Comparing latent vectors across samples helps interpret which dimensions "
        "encode resistance-relevant depth patterns."
)
body(
    "Why this matters: the latent vector is the compressed fingerprint of a sample's "
    "depth landscape. Samples with similar resistance profiles tend to cluster in latent "
    "space. Inspecting which dimensions activate strongly for known positives guides "
    "understanding of what the model has learned."
)

# 4.3 Read-depth track
heading("4.3  Chromosomal read-depth track with HMM segmentation", level=2)

body(
    "The read-depth track shows the raw per-bin coverage alongside the VAE reconstruction "
    "and the HMM copy-number segments for one sample and one chromosome. This is the "
    "primary visual for understanding individual calls."
)
figure(
    f"{SCREENSHOTS}/exp29_depth_ERR10461880.png",
    width_inches=6.0,
    cap="Figure 3. Read-depth track for sample ERR10461880 on chromosome NC_016845.1 "
        "(K. pneumoniae HS11286, ~5.3 Mb). Upper panel: copy ratio (input / VAE reconstruction); "
        "coloured horizontal bars are HMM copy-number segments (grey = CN 1, orange = CN 2, "
        "dark blue = CN 0). Dashed line at copy ratio = 1.0. Lower panel: raw read depth "
        "(green) and VAE reconstruction (red). "
        "The VAE reconstruction closely tracks the smooth chromosomal baseline, filtering "
        "out local noise. A short CN = 2 segment is visible near position ~100 kb (orange), "
        "and two CN = 0 dips near 300 kb (dark blue) correspond to low-coverage regions "
        "likely representing repetitive elements excluded from the core genome."
)
body(
    "Why this matters: the HMM segmentation converts a noisy continuous signal into "
    "interpretable copy-number calls. Visualising the raw depth, reconstruction, and segments "
    "together allows the scientist to verify whether a call reflects genuine biology "
    "(sustained depth change) or a mapping artefact (a single-bin spike). The VAE acts as "
    "a denoising prior: because it was trained on the full cohort, it learns the expected "
    "depth profile and produces a smooth reconstruction that reveals true CNV segments."
)

# 4.4 PCN distributions
heading("4.4  Plasmid copy number distributions", level=2)

body(
    "For each plasmid-borne resistance gene, the histogram below shows the distribution "
    "of PCN values across all 545 samples. PCN is defined as the mean depth at the gene "
    "divided by the sample's chromosomal median depth, giving a copy-number-relative-to-chromosome "
    "estimate."
)
figure(
    f"{SCREENSHOTS}/exp29_pcn_distributions.png",
    width_inches=6.5,
    cap="Figure 4. Plasmid copy number (PCN) distributions for all eight plasmid genes, "
        "experiment 29 (x-axis clipped at PCN < 5 for clarity; blaOXA-48 and blaKPC-2 "
        "have a small number of samples with PCN > 5 indicating high-copy plasmids). "
        "Each panel reports the number of samples called 'present' (CN ≥ 1) in parentheses. "
        "blaCTX-M-14 and blaTEM-1 show zero positive calls because reads from these genes "
        "cross-map to blaCTX-M-15 and chromosomal blaSHV respectively under unmapped-read "
        "remapping — a known alignment limitation. "
        "aac6-Ib-cr positive samples cluster at very low PCN (0.1–0.6), reflecting low "
        "plasmid copy number for this gene family; the absent threshold was lowered to 0.10 "
        "in experiment 29 to recover these samples."
)
body(
    "Why this matters: the PCN distribution for each gene reveals whether the absent/present "
    "threshold is well-placed. A clear bimodal distribution (absent peak at PCN ≈ 0, present "
    "peak at PCN ≥ 0.5) indicates reliable calling. A broad or overlapping distribution — "
    "as seen for aac6-Ib-cr — signals that the gene's plasmid has low copy number in many "
    "carriers, requiring a lower detection threshold and increased vigilance about false "
    "negatives at the boundary."
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. Diagnostics app
# ─────────────────────────────────────────────────────────────────────────────
heading("5. Interactive diagnostics application")

body(
    "The Streamlit diagnostics app (diagnostics/app.py) provides interactive exploration of "
    "experiment outputs. It is started from the repository root with:"
)
p = doc.add_paragraph()
run = p.add_run("    cd diagnostics && streamlit run app.py")
run.font.name = "Courier New"
run.font.size = Pt(10)
doc.add_paragraph()

body("The app has three pages:")

heading("5.1  Monitor", level=2)
body(
    "Displays the training loss curves (reconstruction loss and KL divergence) for a "
    "selected results directory, updating live every 3 seconds. Used to track convergence "
    "during long HPC training runs. When the KL term plateaus and reconstruction loss "
    "stabilises, training is considered complete."
)

heading("5.2  Sample viewer", level=2)
body(
    "The primary diagnostic page. After selecting an experiment, the user can:"
)
bullet("Browse all samples in the cohort via a dropdown or 'Random sample' / "
       "'I'm Feeling Lucky' buttons.")
bullet("Filter by gene call status using the 'Filter by gene calls' expander. "
       "Plasmid genes (blaKPC-2, blaCTX-M-15, blaNDM-1, qnrB1, blaOXA-48, aac6-Ib-cr) "
       "and chromosomal genes (blaSHV, ramR) each have a multiselect for copy-number state "
       "(absent / present / amplified). Conditions are ANDed — for example, selecting "
       "blaKPC-2 = 1 and blaNDM-1 = 1 returns only dual-carbapenemase-carrying samples.")
bullet("Inspect the latent vector bar chart (Figure 2) and PCA scatter plot showing where "
       "the selected sample sits relative to the full cohort.")
bullet("Browse the read-depth track (Figure 3) chromosome by chromosome, with HMM "
       "copy-number segments overlaid.")
bullet("View the gene calls table (chromosomal and plasmid CN per gene) for the "
       "selected sample at the bottom of the page.")
body(
    "Clinical use case: a clinician or microbiologist can select all blaKPC-2-positive "
    "samples, then further filter by blaNDM-1 presence to identify the subset carrying "
    "both carbapenemases — a particularly treatment-resistant combination. The depth track "
    "for each such sample can then be inspected to confirm the calls are supported by "
    "robust signal."
)

heading("5.3  Latent coverage", level=2)
body(
    "Generates a UMAP embedding of the 10-dimensional latent space, with void probes — "
    "synthetic points placed in the empty regions between real samples — coloured by "
    "their distance to the nearest real sample. This visualisation shows where the latent "
    "space is densely populated (well-generalised by the model) versus sparse (under-represented "
    "depth profiles)."
)
body(
    "Hovering over any point in the UMAP reveals the full latent vector in the bar chart "
    "panel on the right, enabling rapid exploration of what genome-wide depth pattern "
    "each region of latent space corresponds to."
)
body(
    "Why this matters: if large regions of latent space are populated only by void probes "
    "(high distance to nearest sample), the model has not seen enough examples of that "
    "depth profile to generalise reliably. Identifying such gaps guides cohort expansion "
    "decisions — for instance, adding more samples from under-represented sequence types "
    "or resistance backgrounds."
)

# ─────────────────────────────────────────────────────────────────────────────
# 6. Limitations and next steps
# ─────────────────────────────────────────────────────────────────────────────
heading("6. Current limitations and next steps")

body("The following limitations are known and are being addressed:")

bullet("blaCTX-M-15 FNR = 0.20 (53 false negatives): reads from CTX-M-15 positive samples "
       "cross-map to chromosomal blaSHV at alignment time. Because the original BAMs were "
       "mapped to HS11286 (which carries blaSHV), those reads are already consumed and do not "
       "appear in the unmapped pool used by the plasmid remap pipeline. This is an irreducible "
       "limitation of the unmapped-read approach for genes with high chromosomal homologues.")
bullet("blaCTX-M-14 and blaTEM-1 — zero positive calls: the short-read sequences of these "
       "variants are too similar to blaCTX-M-15 and chromosomal blaSHV respectively to be "
       "uniquely remapped. Long-read data or targeted amplicon approaches would be needed "
       "to distinguish these at the read level.")
bullet("aac6-Ib-cr FNR = 0.13: ~18 samples with PCN < 0.10 are below the detection limit. "
       "These likely carry the gene on a very low-copy plasmid. Further threshold reduction "
       "risks increasing false positives from background noise.")
bullet("blaSHV chromosomal amplification and ramR deletion — calling not yet evaluated: "
       "Kleborate-derived ground truth is available; integration into the evaluation pipeline "
       "is the next development priority.")
bullet("No sequence-type stratification: the cohort metadata does not yet include ST calls, "
       "preventing per-ST performance analysis which would be clinically informative "
       "(e.g. ST258 is the dominant KPC-carrying lineage).")

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
